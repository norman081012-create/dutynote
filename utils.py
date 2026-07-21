import json
import os
import re
import io
import unicodedata
from copy import deepcopy
from datetime import timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.table import _Row

# ================= 全域設定 =================
tw_tz = timezone(timedelta(hours=8))
TEMPLATE_PATH = "template.docx"
DB_FILE = "handovers.json"

ATTENDING_DOCS_GLOBAL = ["", "鍾偉倫", "張志華", "成毓賢", "劉俊麟", "謝金村", "簡維廷", "唐銘駿", "張維紘"]
ATTENDING_DOCS_FORM = ["未選擇", "鍾偉倫", "張志華", "成毓賢", "劉俊麟", "謝金村", "簡維廷", "唐銘駿", "張維紘"]
DIAG_CHOICES_FORM = ["未選擇", "Schizophrenia", "Bipolar", "Depression", "其他 (請於下方輸入)"]

# ================= 資料處理與排序函數 =================
def load_handovers():
    if os.path.exists(DB_FILE):
        with open(DB_FILE, "r", encoding="utf-8") as f:
            try: return json.load(f)
            except: return []
    return []

def save_handovers(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def parse_his_data(raw_text):
    parsed_stations = {}
    parsed_new = []
    parsed_out = []
    if raw_text:
        for line in raw_text.splitlines():
            line = line.strip()
            if not line: continue
            parts = [p.strip() for p in re.split(r'\t|\s{2,}', line)]
            row_str = "".join(parts).replace(" ", "")
            if "危險評估" in row_str or "自殺顧慮" in row_str: continue
            matched_station = False
            for key_name in ["急診護理站", "二樓護理站", "三樓護理站", "四樓護理站", "五樓護理站", "總人數"]:
                if key_name in row_str and len(parts) >= 4:
                    for idx, p in enumerate(parts):
                        if key_name in p.replace(" ", ""):
                            if idx + 3 < len(parts):
                                parsed_stations[key_name] = parts[idx+1 : idx+4]
                            matched_station = True
                            break
                    if matched_station: break
            if matched_station: continue
            if len(parts) >= 5 and "姓名" not in row_str and "病患" not in row_str:
                if len(parts) >= 7 and ("紅" in row_str or "黃" in row_str or "綠" in row_str or len(parts[6]) < 4):
                    parsed_new.append(parts)
                else:
                    parsed_out.append(parts)
    return parsed_stations, parsed_new, parsed_out

def parse_prn_data(raw_text):
    if not raw_text.strip(): return ""
    doc_map = {}
    for line in raw_text.splitlines():
        parts = [p.strip() for p in line.split('\t')]
        if len(parts) < 6: continue
        name, att, med_full = parts[1], parts[3], parts[5] 
        match = re.search(r'[A-Za-z]+', med_full)
        if not match: continue 
        med = match.group(0)
        if len(med) <= 4: continue
        
        if att not in doc_map: doc_map[att] = {}
        if name not in doc_map[att]: doc_map[att][name] = []
        if med not in doc_map[att][name]: doc_map[att][name].append(med)
    
    output_lines = []
    for doc in ATTENDING_DOCS_GLOBAL:
        if doc and doc in doc_map:
            pt_list = []
            for pt, meds in doc_map[doc].items():
                pt_list.append(f"{pt}{'+'.join(meds)}")
            output_lines.append(f"{doc}：{'，'.join(pt_list)}")
            
    if output_lines:
        return "【PRN 藥物使用】\n" + "\n".join(output_lines)
    return ""

def get_sort_key(h):
    loc = h.get('location', '')
    is_special = h.get('is_special', False)
    if loc == "急診" and is_special: p_loc = 1
    elif loc == "急診" and not is_special: p_loc = 2
    elif is_special: p_loc = 3 
    elif loc == "病房": p_loc = 4
    elif loc == "二樓病房": p_loc = 5
    elif loc == "三樓病房": p_loc = 6
    elif loc == "四樓病房": p_loc = 7
    elif loc == "五樓病房": p_loc = 8
    else: p_loc = 9

    t_str = h.get('time_occurred', '00:00')
    try:
        hrs, mins = map(int, t_str.split(':'))
        total_mins = hrs * 60 + mins
    except: total_mins = 0

    p_time_block = 1 if 480 <= total_mins <= 1439 else 2
    return (p_loc, p_time_block, total_mins)


# ================= Word 排版與輸出函數 =================
def get_unique_cells(row):
    unique_cells = []
    for cell in row.cells:
        if cell not in unique_cells: unique_cells.append(cell)
    return unique_cells

def safe_fill_cell(cell, text, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    if text is None: text = ""
    for p in cell.paragraphs: p.text = "" 
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(str(text).strip())
    run.font.size = Pt(font_size)
    run.bold = False
    p.alignment = align
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def get_text_width(text):
    width = 0
    for char in text:
        if unicodedata.east_asian_width(char) in ('F', 'W', 'A'): width += 2
        else: width += 1
    return width

def visual_smart_chunker(text, max_visual_width=78):
    if not text: return []
    tokens = re.findall(r'[a-zA-Z0-9.\-\_]+|.', text)
    chunks = []
    current_chunk = ""
    current_width = 0
    for token in tokens:
        token_width = get_text_width(token)
        if current_width + token_width > max_visual_width:
            if current_chunk: chunks.append(current_chunk.strip())
            current_chunk = token.lstrip()
            current_width = get_text_width(current_chunk)
        else:
            current_chunk += token
            current_width += token_width
    if current_chunk: chunks.append(current_chunk.strip())
    return chunks

def build_word_and_check_overflow(p_stations, p_new, p_out, lines, selected_date):
    if not os.path.exists(TEMPLATE_PATH): raise FileNotFoundError(f"找不到 {TEMPLATE_PATH}。")
    doc = Document(TEMPLATE_PATH)
    
    roc_year = selected_date.year - 1911
    date_str = f"日期： {roc_year} 年 {selected_date.month:02d} 月 {selected_date.day:02d} 日"

    # 替換日期
    for p in doc.paragraphs:
        txt = p.text.replace(" ", "")
        if "日期" in txt and ("年" in txt or "月" in txt):
            try: p.text = re.sub(r'日期[：:].*日', date_str, p.text)
            except: p.text = date_str
    
    # 填充護理站動態
    for table in doc.tables:
        for row in table.rows:
            u_cells = get_unique_cells(row)
            row_txt = "".join([c.text for c in u_cells]).replace(" ", "")
            matched_st = None
            for kn in ["急診護理站", "二樓護理站", "三樓護理站", "四樓護理站", "五樓護理站", "總人數"]:
                if kn in row_txt: matched_st = kn; break
            if matched_st and matched_st in p_stations:
                for idx, c in enumerate(u_cells):
                    if matched_st in re.sub(r'[\r\n\t]', '', c.text.replace(" ", "").replace(" ", "")) and idx+3 < len(u_cells):
                        safe_fill_cell(u_cells[idx+1], p_stations[matched_st][0], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                        safe_fill_cell(u_cells[idx+2], p_stations[matched_st][1], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                        safe_fill_cell(u_cells[idx+3], p_stations[matched_st][2], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 處理出入院病人 (刪除空白行)
    for table in doc.tables:
        blank_new_rows, blank_out_rows, section = [], [], None
        name_col_new, name_col_out = 0, 0
        for row in table.rows:
            u_cells = get_unique_cells(row)
            row_txt = "".join(c.text for c in u_cells).replace(" ", "")
            if "姓名" in row_txt and "病歷" in row_txt:
                if "燈號" in row_txt or "強制" in row_txt: 
                    section, name_col_new = "new", next((i for i, c in enumerate(u_cells) if "姓名" in c.text.replace(" ", "")), 0)
                else: 
                    section, name_col_out = "out", next((i for i, c in enumerate(u_cells) if "姓名" in c.text.replace(" ", "")), 0)
            elif "出院病人" in row_txt or "危險評估" in row_txt: section = None
            elif section == "new" and name_col_new < len(u_cells):
                if re.sub(r'[\r\n\t\s_0]', '', u_cells[name_col_new].text) == "": blank_new_rows.append((row, name_col_new))
            elif section == "out" and name_col_out < len(u_cells):
                if re.sub(r'[\r\n\t\s_0]', '', u_cells[name_col_out].text) == "": blank_out_rows.append((row, name_col_out))

        while len(p_new) > len(blank_new_rows) and blank_new_rows:
            last_row, col = blank_new_rows[-1]
            new_tr = deepcopy(last_row._tr)
            last_row._tr.addnext(new_tr)
            blank_new_rows.append((_Row(new_tr, last_row._parent), col))
            
        while len(p_out) > len(blank_out_rows) and blank_out_rows:
            last_row, col = blank_out_rows[-1]
            new_tr = deepcopy(last_row._tr)
            last_row._tr.addnext(new_tr)
            blank_out_rows.append((_Row(new_tr, last_row._parent), col))

        for i, (row, col_idx) in enumerate(blank_new_rows):
            if i < len(p_new):
                u_cells = get_unique_cells(row)
                for k in range(min(len(p_new[i]), len(u_cells))):
                    tc = col_idx + k if k < 6 else col_idx + k + 1
                    if tc < len(u_cells): safe_fill_cell(u_cells[tc], p_new[i][k], font_size=10)
            else:
                try: row._element.getparent().remove(row._element)
                except: pass
                
        for i, (row, col_idx) in enumerate(blank_out_rows):
            if i < len(p_out):
                u_cells = get_unique_cells(row)
                for k in range(min(len(p_out[i]), len(u_cells))):
                    if col_idx + k < len(u_cells): safe_fill_cell(u_cells[col_idx + k], p_out[i][k], font_size=10)
            else:
                try: row._element.getparent().remove(row._element)
                except: pass

    # ================= 核心邏輯：動態分頁與完美保存底部新版簽章 =================
    all_chunks_to_fill = []
    for line in lines:
        if line == "": all_chunks_to_fill.append("")
        else: all_chunks_to_fill.extend(visual_smart_chunker(line, max_visual_width=78))

    target_table, header_row_idx, discuss_row_idx = None, -1, -1
    for table in doc.tables:
        for idx, row in enumerate(table.rows):
            u_cells = get_unique_cells(row)
            if not u_cells: continue
            row_txt = u_cells[0].text.replace(" ", "")
            if "病房特殊狀況及處理" in row_txt: header_row_idx = idx
            if "討論與講評" in row_txt: discuss_row_idx = idx
            if header_row_idx != -1 and discuss_row_idx != -1: target_table = table; break
        if target_table: break

    is_overflow = False 
    
    if target_table and header_row_idx != -1 and discuss_row_idx != -1:
        # 1. 提取並備份「討論與講評」及底部所有帶有勾選框的簽章區塊 XML
        # 因為我們是直接備份整個區塊 XML，所以不論您模板有4行還是6行空白，都會被一字不漏保留
        discuss_rows_xml = []
        for r_idx in range(discuss_row_idx, len(target_table.rows)):
            discuss_rows_xml.append(deepcopy(target_table.rows[r_idx]._element))
        
        # 2. 從原表格中刪除這些底部區塊，將表格變成純粹的資料填寫區
        rows_to_remove = [r._element for r in target_table.rows[discuss_row_idx:]]
        for r_elem in rows_to_remove:
            target_table._element.remove(r_elem)

        # 3. 計算單頁可容納的交班事項行數
        start_row_idx = header_row_idx + 1
        capacity = discuss_row_idx - start_row_idx
        
        chunks_on_first_page = all_chunks_to_fill[:capacity]
        chunks_overflow = all_chunks_to_fill[capacity:]

        # 4. 填寫第一頁內容
        for i, chunk_text in enumerate(chunks_on_first_page):
            cell = get_unique_cells(target_table.rows[start_row_idx + i])[0]
            safe_fill_cell(cell, chunk_text, font_size=12)

        final_table_xml = target_table._element
        
        # 5. 若溢出，複製原表格架構，建立新頁面
        if chunks_overflow:
            is_overflow = True
            chunk_groups = [chunks_overflow[i:i + 15] for i in range(0, len(chunks_overflow), 15)]

            for group in chunk_groups:
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
                
                # 複製一個包含原始表格所有欄寬設定的全新表格
                new_table_xml = deepcopy(target_table._element)
                all_trs = new_table_xml.xpath('.//w:tr')
                for tr in all_trs[header_row_idx + 1:]:
                    new_table_xml.remove(tr)

                source_row_xml = deepcopy(target_table.rows[start_row_idx]._element)
                for chunk_text in group:
                    new_row = deepcopy(source_row_xml)
                    new_table_xml.append(new_row)
                    row_obj = _Row(new_row, target_table)
                    cell = get_unique_cells(row_obj)[0]
                    safe_fill_cell(cell, chunk_text, font_size=12)

                doc._body._element.append(new_table_xml)
                final_table_xml = new_table_xml

        # 6. 將剛剛備份的「討論與講評與新版簽章區塊」完美拼貼至最後一個表格尾端
        for r_xml in discuss_rows_xml:
            final_table_xml.append(deepcopy(r_xml))

    stream = io.BytesIO(); doc.save(stream); stream.seek(0)
    return stream, is_overflow
