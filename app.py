import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import _Row
from copy import deepcopy
import unicodedata
import io
import json
import os
import re
import datetime
from datetime import timezone, timedelta

# 設定台灣時區 (UTC+8)
tw_tz = timezone(timedelta(hours=8))

st.set_page_config(page_title="值班日誌自動生成器", layout="wide")

TEMPLATE_PATH = "template.docx"
DB_FILE = "handovers.json"

# UI 專用的選項名單 (含「太忙了沒時間問」)
ATTENDING_DOCS_GLOBAL = ["", "鍾偉倫", "張志華", "成毓賢", "劉俊麟", "謝金村", "唐銘駿", "吳騂", "張維紘"]
ATTENDING_DOCS_FORM = ["太忙了沒時間問", "鍾偉倫", "張志華", "成毓賢", "劉俊麟", "謝金村", "唐銘駿", "吳騂", "張維紘"]
DIAG_CHOICES_FORM = ["太忙了沒時間問", "Schizophrenia", "bipolar", "depression", "其他 (請於下方輸入)"]

# 年齡選單 (往上 49~1，預設太忙，往下 50~110)
age_options = [str(i) for i in range(1, 50)] + ["太忙了沒時間問"] + [str(i) for i in range(50, 111)]
default_age_idx = age_options.index("太忙了沒時間問")

# --- CSS 樣式注入 (輸入框置中 & 消除多餘空白) ---
st.markdown("""
<style>
/* 讓文字輸入框置中 */
div[data-baseweb="input"] input {
    text-align: center !important;
}
/* 讓下拉選單文字置中 */
div[data-baseweb="select"] div {
    text-align: center !important;
    justify-content: center !important;
}
/* 消除提示框與下方標題之間的空白 */
div[data-testid="stAlert"] {
    margin-bottom: 0px !important;
    padding-top: 10px !important;
    padding-bottom: 10px !important;
}
h2 {
    padding-top: 0.5rem !important;
}
</style>
""", unsafe_allow_html=True)

# ================= 資料處理函數 =================
def load_handovers():
    if os.path.exists(DB_FILE):
        with open(DB_FILE, "r", encoding="utf-8") as f:
            try: return json.load(f)
            except: return []
    return []

def save_handovers(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def parse_his_data(raw_text):
    parsed_stations = {}
    parsed_new = []
    parsed_out = []
    if raw_text:
        for line in raw_text.splitlines():
            line = line.strip()
            if not line: continue
            parts = [p.strip() for p in re.split(r'\t|\s{2,}', line)]
            row_str = "".join(parts).replace(" ", "")
            if "危險評估" in row_str or "自殺顧慮" in row_str: continue
            matched_station = False
            for key_name in ["急診護理站", "二樓護理站", "三樓護理站", "四樓護理站", "五樓護理站", "總人數"]:
                if key_name in row_str and len(parts) >= 4:
                    for idx, p in enumerate(parts):
                        if key_name in p.replace(" ", ""):
                            if idx + 3 < len(parts):
                                parsed_stations[key_name] = parts[idx+1 : idx+4]
                            matched_station = True
                            break
                    if matched_station: break
            if matched_station: continue
            if len(parts) >= 5 and "姓名" not in row_str and "病患" not in row_str:
                if len(parts) >= 7 and ("紅" in row_str or "黃" in row_str or "綠" in row_str or len(parts[6]) < 4):
                    parsed_new.append(parts)
                else:
                    parsed_out.append(parts)
    return parsed_stations, parsed_new, parsed_out

def parse_prn_data(raw_text):
    """解析 PRN 藥物文字並按醫師分類"""
    if not raw_text.strip(): return ""
    doc_map = {}
    for line in raw_text.splitlines():
        parts = [p.strip() for p in line.split('\t')]
        if len(parts) < 6: continue
        name, att, med_full = parts[1], parts[3], parts[5]
        # 藥名簡化：取第一個單字
        med = med_full.split(' ')[0].split('(')[0].strip()
        
        if att not in doc_map: doc_map[att] = {}
        if name not in doc_map[att]: doc_map[att][name] = []
        if med not in doc_map[att][name]: doc_map[att][name].append(med)
    
    output_lines = []
    for doc in ATTENDING_DOCS_GLOBAL:
        if doc and doc in doc_map:
            pt_list = []
            for pt, meds in doc_map[doc].items():
                pt_list.append(f"{pt}{'+'.join(meds)}")
            output_lines.append(f"{doc}：{'，'.join(pt_list)}")
    
    if output_lines:
        return "【PRN 藥物使用】\n" + "\n".join(output_lines)
    return ""

# ================= 表單狀態與全域設定初始化 =================
if 'handovers' not in st.session_state: st.session_state.handovers = load_handovers()
if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0

now_tw = datetime.datetime.now(tw_tz)
if "f_duty_date" not in st.session_state: st.session_state.f_duty_date = now_tw.date()
if "f_duty_doc" not in st.session_state: st.session_state.f_duty_doc = ""
    
if "f_loc" not in st.session_state:
    st.session_state.update({
        "f_loc": "病房", "f_name": "", "f_age": "太忙了沒時間問", "f_gen": "",
        "f_med": "", "f_hist": "", "f_time": now_tw.time(),
        "f_doc": "太忙了沒時間問", "f_diag_c": "太忙了沒時間問", "f_diag_m": "", "f_content": "",
        "f_special": False, "add_error": False
    })

# ================= 排序與 Callback =================
def get_sort_key(h):
    loc = h.get('location', '')
    is_special = h.get('is_special', False)

    if loc == "急診" and is_special: p_loc = 1
    elif loc == "急診" and not is_special: p_loc = 2
    elif is_special: p_loc = 3 
    elif loc == "病房": p_loc = 4
    elif loc == "二樓病房": p_loc = 5
    elif loc == "三樓病房": p_loc = 6
    elif loc == "四樓病房": p_loc = 7
    elif loc == "五樓病房": p_loc = 8
    else: p_loc = 9

    t_str = h.get('time_occurred', '00:00')
    try:
        hrs, mins = map(int, t_str.split(':'))
        total_mins = hrs * 60 + mins
    except: total_mins = 0

    p_time_block = 1 if 480 <= total_mins <= 1439 else 2
    return (p_loc, p_time_block, total_mins)

def clear_form():
    st.session_state.update({
        "f_loc": "病房", "f_name": "", "f_age": "太忙了沒時間問", "f_gen": "",
        "f_med": "", "f_hist": "", "f_time": datetime.datetime.now(tw_tz).time(),
        "f_doc": "太忙了沒時間問", "f_diag_c": "太忙了沒時間問", "f_diag_m": "", 
        "f_content": "", "f_special": False, "add_error": False
    })

def load_form(h):
    st.session_state.f_loc = h.get("location", "病房")
    st.session_state.f_name = h.get("name", "")
    age = h.get("age", "")
    st.session_state.f_age = "太忙了沒時間問" if age == "" else age
    st.session_state.f_gen = h.get("gender", "")
    st.session_state.f_med = h.get("med_record", "")
    st.session_state.f_hist = h.get("history", "")
    try: st.session_state.f_time = datetime.datetime.strptime(h.get("time_occurred", "00:00"), "%H:%M").time()
    except: st.session_state.f_time = datetime.datetime.now(tw_tz).time()
        
    doc = h.get("attending_doc", "")
    st.session_state.f_doc = "太忙了沒時間問" if doc == "" else doc
    
    diag = h.get("diagnosis", "")
    if diag in ["Schizophrenia", "bipolar", "depression"]:
        st.session_state.f_diag_c = diag
        st.session_state.f_diag_m = ""
    elif diag == "":
        st.session_state.f_diag_c = "太忙了沒時間問"
        st.session_state.f_diag_m = ""
    else:
        st.session_state.f_diag_c = "其他 (請於下方輸入)"
        st.session_state.f_diag_m = diag
        
    st.session_state.f_content = h.get("content", "")
    st.session_state.f_special = h.get("is_special", False)

def cb_refresh():
    st.session_state.handovers = []
    save_handovers([])
    clear_form()
    st.session_state.uploader_key += 1
    st.session_state.f_duty_date = datetime.datetime.now(tw_tz).date()
    st.session_state.f_duty_doc = ""

def cb_add():
    if not st.session_state.f_name or not st.session_state.f_content:
        st.session_state.add_error = True
    else:
        st.session_state.add_error = False
        diag_c_val = "" if st.session_state.f_diag_c == "太忙了沒時間問" else st.session_state.f_diag_c
        diag_final = st.session_state.f_diag_m if not diag_c_val or diag_c_val == "其他 (請於下方輸入)" else diag_c_val
        age_val = "" if st.session_state.f_age == "太忙了沒時間問" else st.session_state.f_age
        doc_val = "" if st.session_state.f_doc == "太忙了沒時間問" else st.session_state.f_doc

        st.session_state.handovers.append({
            "location": st.session_state.f_loc, "name": st.session_state.f_name, 
            "age": age_val, "gender": st.session_state.f_gen,
            "med_record": st.session_state.f_med, "attending_doc": doc_val,
            "time_occurred": st.session_state.f_time.strftime("%H:%M"), "content": st.session_state.f_content,
            "diagnosis": diag_final, "history": st.session_state.f_hist,
            "is_er": (st.session_state.f_loc == "急診"),
            "is_special": st.session_state.f_special
        })
        save_handovers(st.session_state.handovers)
        clear_form()

def cb_edit(idx, h):
    load_form(h)
    st.session_state.handovers.pop(idx)
    save_handovers(st.session_state.handovers)

def cb_delete(idx):
    st.session_state.handovers.pop(idx)
    save_handovers(st.session_state.handovers)


# ================= UI 畫面區 =================
st.title("🏥 醫師病房值班日誌自動生成器")

col_warn, col_btn = st.columns([8, 2], vertical_alignment="center")
with col_warn:
    st.info("⚠️ **溫馨提示：** 新值班醫師接班時，請務必點擊右方的「🔄 刷新並清空所有資料」，否則將會讀取到前一位醫師的設定檔與暫存資料喔！")
with col_btn:
    st.button("🔄 刷新並清空所有資料", type="secondary", use_container_width=True, on_click=cb_refresh)

st.header("1. 貼上系統匯出資料")
c_date, c_his, c_prn = st.columns([2, 4, 4])

with c_date:
    st.date_input("📅 選擇值班日期", key="f_duty_date")
    st.selectbox("👨‍⚕️ 選擇值班醫師", ATTENDING_DOCS_GLOBAL, key="f_duty_doc")

with c_his:
    raw_his = st.text_area("📝 貼上 HIS 內容 (人數/出入院)", height=150, key=f"his_{st.session_state.uploader_key}")
    parsed_stations, parsed_new, parsed_out = parse_his_data(raw_his)

with c_prn:
    raw_prn = st.text_area("💊 貼上 PRN 藥物清單 (選填)", height=150, key=f"prn_{st.session_state.uploader_key}")
    prn_summary = parse_prn_data(raw_prn)

# ================= 區塊 2：交班事項登錄表單 =================
st.header("2. 交班事項登錄")
c1, c2 = st.columns(2)
with c1:
    st.selectbox("單位/病房 (預設此)", ["病房", "急診", "二樓病房", "三樓病房", "四樓病房", "五樓病房"], key="f_loc")
    st.text_input("病人姓名 (必填)", key="f_name")
    st.selectbox("年紀", age_options, index=default_age_idx, key="f_age")
    st.selectbox("性別", ["", "男", "女"], key="f_gen")
    st.text_input("病歷號", key="f_med")
    st.text_area("內外科病史輸入", height=60, key="f_hist")
    
with c2:
    st.time_input("狀況發生時間", key="f_time")
    st.selectbox("主治醫師", ATTENDING_DOCS_FORM, key="f_doc")
    st.selectbox("診斷快速選項", DIAG_CHOICES_FORM, key="f_diag_c")
    st.text_input("手動輸入診斷 (若選其他)", key="f_diag_m")
    st.checkbox("🚨 特別交班", key="f_special")
    
st.text_area("交班內容 (必填)", key="f_content")

btn_col1, btn_col2, btn_col3 = st.columns([2, 1, 1])
with btn_col1:
    st.button("✅ 確認新增交班", type="primary", use_container_width=True, on_click=cb_add)
    if st.session_state.add_error: st.error("「姓名」與「內容」為必填！")
with btn_col2:
    st.button("🔄 重新輸入", use_container_width=True, on_click=clear_form)

# ================= 區塊 3：已登錄交班預覽 =================
st.header("3. 已登錄交班事項")
if st.session_state.handovers:
    sorted_view = sorted(st.session_state.handovers, key=get_sort_key)
    for h in sorted_view:
        idx = st.session_state.handovers.index(h)
        h_age_disp = h['age'] if h.get('age') else "?"
        h_gen_disp = f"{h['gender']}性" if h.get('gender') else ""
        sp_tag = " [🚨特別交班]" if h.get('is_special') else ""
        
        with st.expander(f"[{h['location']}] {h['name']} ({h_age_disp}歲{h_gen_disp}) - {h['time_occurred']}{sp_tag}"):
            h_diag_disp = h['diagnosis'] if h.get('diagnosis') else "??"
            st.write(f"主治：{h['attending_doc']} | 病史：{h['history']} | 診斷：{h_diag_disp}")
            st.write(f"內容：{h['content']}")
            
            c_edit, c_del, c_empty = st.columns([1.5, 1.5, 7])
            with c_edit: st.button(f"✏️ 修改 {h['name']}", key=f"edit_{idx}", on_click=cb_edit, args=(idx, h))
            with c_del: st.button(f"🗑️ 刪除 {h['name']}", key=f"del_{idx}", on_click=cb_delete, args=(idx,))

# ================= 工具與輸出 =================
def get_unique_cells(row):
    unique_cells = []
    for cell in row.cells:
        if cell not in unique_cells: unique_cells.append(cell)
    return unique_cells

def safe_fill_cell(cell, text, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    if text is None: text = ""
    for p in cell.paragraphs: p.text = "" 
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(str(text).strip())
    run.font.size = Pt(font_size)
    run.bold = False
    p.alignment = align
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def get_text_width(text):
    width = 0
    for char in text:
        if unicodedata.east_asian_width(char) in ('F', 'W', 'A'): width += 2
        else: width += 1
    return width

def visual_smart_chunker(text, max_visual_width=78):
    if not text: return []
    tokens = re.findall(r'[a-zA-Z0-9.\-\_]+|.', text)
    chunks = []
    current_chunk = ""
    current_width = 0
    for token in tokens:
        token_width = get_text_width(token)
        if current_width + token_width > max_visual_width:
            if current_chunk: chunks.append(current_chunk.strip())
            current_chunk = token.lstrip()
            current_width = get_text_width(current_chunk)
        else:
            current_chunk += token
            current_width += token_width
    if current_chunk: chunks.append(current_chunk.strip())
    return chunks

st.header("4. 預覽與輸出")

# 生成最終預覽文字
preview_lines = []
sorted_h = sorted(st.session_state.handovers, key=get_sort_key)
for h in sorted_h:
    h_loc = h.get('location', '病房')
    h_name = h.get('name', '').strip()
    h_age = h.get('age', '').strip()
    h_gen = h.get('gender', '').strip()
    h_med = h.get('med_record', '').strip()
    h_att = h.get('attending_doc', '').strip()
    h_diag = h.get('diagnosis', '').strip()
    h_his = h.get('history', '').strip()
    h_time = h.get('time_occurred', '').strip()
    h_content = h.get('content', '').replace('\n', ' ').strip()

    h_age_display = h_age if h_age else "?"
    h_gen_display = f"{h_gen}性" if h_gen else ""
    age_gen_part = f"，{h_age_display}歲{h_gen_display}"
    med_part = f"病歷號:{h_med} " if h_med else ""
    pt_part = f"({h_loc}){med_part}姓名:{h_name}{age_gen_part}"
    
    ward_tag = f"({h_loc[0:2]})" if h_loc not in ["急診", "病房"] else ""
    doc_part = f"{h_att}醫師{ward_tag}病人" if h_att else ""
    his_part = f"內外科病史:{h_his}" if h_his else ""
    if not h_diag: h_diag = "??"
    diag_part = f"診斷:{h_diag}"
    time_part = f"約{h_time}時" if h_time else ""
    
    diag_time = ""
    if diag_part and time_part: diag_time = f"{diag_part} {time_part}"
    elif diag_part: diag_time = diag_part
    elif time_part: diag_time = time_part
        
    components = [c for c in [pt_part, doc_part, his_part, diag_time, h_content] if c.strip()]
    preview_lines.append("，".join(components))

# 附加 PRN 藥物
if prn_summary:
    preview_lines.append("")
    preview_lines.extend(prn_summary.splitlines())

if preview_lines:
    with st.expander("👀 點擊展開：最終交班文字預覽 (與 Word 輸出內容相同)", expanded=True):
        st.text_area("即將寫入 Word 的文字：", value="\n\n".join(preview_lines), height=250, disabled=True)

# 生成 Word 檔案
def build_word_and_check_overflow(p_stations, p_new, p_out, lines, selected_date, selected_doc):
    if not os.path.exists(TEMPLATE_PATH): raise FileNotFoundError(f"找不到 {TEMPLATE_PATH}。")
    doc = Document(TEMPLATE_PATH)
    
    roc_year = selected_date.year - 1911
    date_str = f"日期： {roc_year} 年 {selected_date.month:02d} 月 {selected_date.day:02d} 日"
    
    def apply_signature(p_element, doc_name):
        p_element.text = "" 
        run_label = p_element.add_run("值班醫師：")
        if doc_name:
            run_name = p_element.add_run(f"  {doc_name}")
            run_name.font.size = Pt(16)
            run_name.bold = True
            run_name.font.name = '標楷體'
            rPr_name = run_name._element.get_or_add_rPr()
            rPr_name.get_or_add_rFonts().set(qn('w:eastAsia'), '標楷體')
        p_element.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in doc.paragraphs:
        txt = p.text.replace(" ", "")
        if "日期" in txt and ("年" in txt or "月" in txt): p.text = date_str
        elif "值班醫師" in txt: apply_signature(p, selected_doc)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "值班醫師" in p.text.replace(" ", ""): apply_signature(p, selected_doc)
    
    for table in doc.tables:
        for row in table.rows:
            u_cells = get_unique_cells(row)
            row_txt = "".join([c.text for c in u_cells]).replace(" ", "")
            matched_st = None
            for kn in ["急診護理站", "二樓護理站", "三樓護理站", "四樓護理站", "五樓護理站", "總人數"]:
                if kn in row_txt: matched_st = kn; break
            if matched_st and matched_st in p_stations:
                for idx, c in enumerate(u_cells):
                    if matched_st in re.sub(r'[\r\n\t]', '', c.text.replace(" ", "").replace("　", "")) and idx+3 < len(u_cells):
                        safe_fill_cell(u_cells[idx+1], p_stations[matched_st][0], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                        safe_fill_cell(u_cells[idx+2], p_stations[matched_st][1], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                        safe_fill_cell(u_cells[idx+3], p_stations[matched_st][2], font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for table in doc.tables:
        blank_new_rows, blank_out_rows, section = [], [], None
        name_col_new, name_col_out = 0, 0
        for row in table.rows:
            u_cells = get_unique_cells(row)
            row_txt = "".join(c.text for c in u_cells).replace(" ", "")
            if "姓名" in row_txt and "病歷" in row_txt:
                if "燈號" in row_txt or "強制" in row_txt: 
                    section, name_col_new = "new", next((i for i, c in enumerate(u_cells) if "姓名" in c.text.replace(" ", "")), 0)
                else: 
                    section, name_col_out = "out", next((i for i, c in enumerate(u_cells) if "姓名" in c.text.replace(" ", "")), 0)
            elif "出院病人" in row_txt or "危險評估" in row_txt: section = None
            elif section == "new" and name_col_new < len(u_cells):
                if re.sub(r'[\r\n\t\s_0]', '', u_cells[name_col_new].text) == "": blank_new_rows.append((row, name_col_new))
            elif section == "out" and name_col_out < len(u_cells):
                if re.sub(r'[\r\n\t\s_0]', '', u_cells[name_col_out].text) == "": blank_out_rows.append((row, name_col_out))

        while len(p_new) > len(blank_new_rows) and blank_new_rows:
            last_row, col = blank_new_rows[-1]
            new_tr = deepcopy(last_row._tr)
            last_row._tr.addnext(new_tr)
            blank_new_rows.append((_Row(new_tr, last_row._parent), col))
            
        while len(p_out) > len(blank_out_rows) and blank_out_rows:
            last_row, col = blank_out_rows[-1]
            new_tr = deepcopy(last_row._tr)
            last_row._tr.addnext(new_tr)
            blank_out_rows.append((_Row(new_tr, last_row._parent), col))

        for i, (row, col_idx) in enumerate(blank_new_rows):
            if i < len(p_new):
                u_cells = get_unique_cells(row)
                for k in range(min(len(p_new[i]), len(u_cells))):
                    tc = col_idx + k if k < 6 else col_idx + k + 1
                    if tc < len(u_cells): safe_fill_cell(u_cells[tc], p_new[i][k], font_size=10)
            else:
                try: row._element.getparent().remove(row._element)
                except: pass
                
        for i, (row, col_idx) in enumerate(blank_out_rows):
            if i < len(p_out):
                u_cells = get_unique_cells(row)
                for k in range(min(len(p_out[i]), len(u_cells))):
                    if col_idx + k < len(u_cells): safe_fill_cell(u_cells[col_idx + k], p_out[i][k], font_size=10)
            else:
                try: row._element.getparent().remove(row._element)
                except: pass

    for table in doc.tables:
        header_row_idx = -1
        for i, row in enumerate(table.rows):
            row_txt = "".join([c.text for c in get_unique_cells(row)]).replace(" ", "")
            if "自殺顧慮" in row_txt and "哽塞顧慮" in row_txt: header_row_idx = i; break
        if header_row_idx != -1:
            while len(table.rows) > header_row_idx + 5:
                row_to_del = table.rows[-1]
                row_to_del._element.getparent().remove(row_to_del._element)
            break

    # 計算交班事項字串，進行溢出偵測
    all_chunks_to_fill = []
    for line in lines:
        if line == "": all_chunks_to_fill.append("")
        else: all_chunks_to_fill.extend(visual_smart_chunker(line, max_visual_width=78))

    target_table, start_row_idx, discuss_row_idx = None, -1, -1
    for table in doc.tables:
        for idx, row in enumerate(table.rows):
            u_cells = get_unique_cells(row)
            if not u_cells: continue
            row_txt = u_cells[0].text.replace(" ", "")
            if "病房特殊狀況及處理" in row_txt: target_table, start_row_idx = table, idx + 1 
            if "討論與講評" in row_txt and start_row_idx != -1: discuss_row_idx = idx; break
        if target_table and discuss_row_idx != -1: break

    # 判斷是否溢出 (所需行數 > 剩餘空白行數)
    is_overflow = False
    if target_table and start_row_idx != -1 and discuss_row_idx != -1:
        capacity = discuss_row_idx - start_row_idx
        if len(all_chunks_to_fill) > capacity:
            is_overflow = True
            
        current_row_idx = start_row_idx
        for chunk_text in all_chunks_to_fill:
            if current_row_idx < discuss_row_idx:
                target_cell = get_unique_cells(target_table.rows[current_row_idx])[0]
                safe_fill_cell(target_cell, chunk_text, font_size=12)
                current_row_idx += 1
            else: # 如果硬擠(動態擴展)，可能會擠壓後方版面，但先容許Word自動換頁
                ref_row = target_table.rows[discuss_row_idx]
                blank_tr = deepcopy(target_table.rows[discuss_row_idx - 1]._tr)
                ref_row._tr.addprevious(blank_tr)
                discuss_row_idx += 1
                target_cell = get_unique_cells(target_table.rows[current_row_idx])[0]
                safe_fill_cell(target_cell, chunk_text, font_size=12)
                current_row_idx += 1
                
        while current_row_idx < discuss_row_idx:
            target_cell = get_unique_cells(target_table.rows[current_row_idx])[0]
            safe_fill_cell(target_cell, "", font_size=12)
            current_row_idx += 1

    stream = io.BytesIO(); doc.save(stream); stream.seek(0)
    return stream, is_overflow

if st.button("🚀 生成下載 Word", type="primary"):
    try:
        f_stream, overflow = build_word_and_check_overflow(
            parsed_stations, parsed_new, parsed_out, 
            preview_lines, 
            st.session_state.f_duty_date, 
            st.session_state.f_duty_doc
        )
        if overflow:
            st.error("⚠️ 交班內容過多（已超過 Word 表格行數），若版面跑位請考慮簡化內容。")
        else:
            st.success("✅ 檔案已更新並備妥！")
            
        st.download_button("📥 點擊下載", f_stream, f"值班日誌_{st.session_state.f_duty_date.strftime('%Y%m%d')}.docx")
    except Exception as e:
        st.error(f"錯誤: {e}")
